import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = 'docs/DriveLegal_Software_Architecture.docx'
img_path = 'public/test-image.jpg'

try:
    doc = Document(doc_path)
    
    first_p = doc.paragraphs[0]
    
    # 1. Insert links at the top
    p_links_title = first_p.insert_paragraph_before('')
    run_title = p_links_title.add_run('🔗 LIVE DEMO & SOURCE CODE')
    run_title.bold = True
    
    first_p.insert_paragraph_before('Live Web Application: https://drive-legal-iit-m.vercel.app/')
    first_p.insert_paragraph_before('GitHub Repository: https://github.com/Chaitanya-970/Drive-Legal---IIT-M')
    
    first_p.insert_paragraph_before('') # Blank line
    
    # 2. Insert AI Testing Note and Image
    p_note = first_p.insert_paragraph_before('')
    run_note = p_note.add_run('👨‍⚖️ Note for Judges: You can test the Edge AI Vehicle Detection functioning by uploading the image below to the Image Scanner! It will accurately detect the vehicle in the image.')
    run_note.bold = True
    
    if os.path.exists(img_path):
        p_img = first_p.insert_paragraph_before('')
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(4.5))
    
    first_p.insert_paragraph_before('') # Blank line
    
    doc.save(doc_path)
    print("Successfully updated docs/DriveLegal_Software_Architecture.docx")
except Exception as e:
    print(f"Error updating docx: {e}")
